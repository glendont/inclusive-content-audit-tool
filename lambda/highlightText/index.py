import json
import boto3
import io
import urllib.parse

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_BREAK
import re


def highlight_banned_words(paragraphs, banned_words, doc):
    # splits the document by paragraphs and runs it through a for loop
    banned_words_used = []
    for paragraph in paragraphs:
        para_text = paragraph.text
        
        #splits the paragraph text into individual elements within a list
        para_text_split = re.findall(r"[\w']+|[.,!?;()-_]", para_text)

        para_banned_words = []
        para_clean = []
        
        for i in range(len(para_text_split)):
            current_word = para_text_split[i]
            if current_word.lower() not in banned_words:
                para_clean.append(current_word)
            else:
                if current_word.lower() not in banned_words_used:
                    banned_words_used.append(current_word.lower())
                para_banned_words.append(current_word)
                changed_word = current_word.replace(current_word, 'replace_this_text')
                para_clean.append(changed_word)
                run = True

        
        para_joined = " ".join(para_clean)
        
        substrings = para_joined.split('replace_this_text')
        
        paragraph = doc.add_paragraph()
        
        counter = 0
        for substring in substrings[:-1]:
            paragraph.add_run(substring)
            font = paragraph.add_run(para_banned_words[counter]).font
            font.highlight_color = WD_COLOR_INDEX.YELLOW
            counter += 1
        paragraph.add_run(substrings[-1])
        paragraph.add_run().add_break(WD_BREAK.LINE)
        
    print(banned_words_used)
    return run, banned_words_used, paragraph, doc



    
def create_banned_words_table(doc, banned_words, banned_words_used):
    para = doc.add_paragraph()
    para.add_run("Biased Words Used")
    para.add_run().add_break(WD_BREAK.LINE)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('You Used The Following Words...').bold = True
    hdr_cells[1].paragraphs[0].add_run('You May Want To Consider Using...').bold = True

    for word in banned_words_used:
        row_cells = table.add_row().cells
        row_cells[0].text = str(word)
        for i in range(len(banned_words[word])):
            if i == (len(banned_words[word])-1):
                run = row_cells[1].paragraphs[0].add_run(banned_words[word][i])
            else:
                run = row_cells[1].paragraphs[0].add_run(banned_words[word][i] + ', ')

def lambda_handler(event, context):
    

    banned_words = {
        'abort': ['stop'],
        'blacklist':['deny list'],
        'execute':['start','run'],
        'hang': ['stop responding'],
        'kill':['end', 'stop'],
        'master':['primary', 'main', 'leader'],
        'slave':['replica', 'secondary', 'standby'],
        'whitelist':['allow list'],
    }
    
        
    s3 = boto3.client('s3')
    # TODO implement
    
    if event:
        print('Event:', event)
        file_obj = event['Records'][0]
        key = urllib.parse.unquote_plus(event['Records'][0]['s3']['object']['key'])
        print(key)
        file_name = key.split('/')[-1]
        file_name_no_suffix = file_name.split('.')[-2]
        bucket = event['Records'][0]['s3']['bucket']['name']
        print('Bucket:', bucket)
        print('Filename:', file_name)
        response = s3.get_object(Bucket=bucket, Key=key)
        
        file = io.BytesIO(response['Body'].read())
        doc = Document(file)
        
        ## get template document
        response = s3.get_object(Bucket="inclusive-content-audit-tool", Key="template/Template.docx")
        file = io.BytesIO(response['Body'].read())
        template_doc = Document(file)
        
        paragraphs = doc.paragraphs
        run = False
        run, banned_words_used, paragraph, doc = highlight_banned_words(paragraphs, banned_words, template_doc)
        
        if run:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            
            result_page_para = doc.add_paragraph()
            result_page_header = result_page_para.add_run('Inclusive Content Test Result')
            result_page_header.bold = True
            result_page_header.add_break(WD_BREAK.LINE)
       
    
            og_doc_para = doc.add_paragraph()
            og_doc_para.add_run('Original Document Name: '+ file_name)
            og_doc_para.add_run().add_break(WD_BREAK.LINE)
            
            words_used_para = doc.add_paragraph()
            total_words_used = len(banned_words_used)
            words_used = words_used_para.add_run(str(total_words_used))
            words_used.bold = True
            words_used.underline = True
            words_used_para.add_run(' words were flagged out for reconsideration')
            words_used_para.add_run().add_break(WD_BREAK.LINE)
            
            if len(banned_words_used) >= 1:
                print('making table')
                create_banned_words_table(doc, banned_words, banned_words_used)
            save_name = f"output/{file_name_no_suffix} - Scanned.docx"
        else:
            doc = Document()
            result_page_para = doc.add_paragraph()
            result_page_header = result_page_para.add_run('Inclusive Content Test Result')
            result_page_header.bold = True
            
            og_doc_para = doc.add_paragraph()
            og_doc_para.add_run('Original Document Name: '+ file_name)
            
            words_used_para = doc.add_paragraph()
            words_used_para.add_run().add_break(WD_BREAK.LINE)
            words_used_para.add_run('Congratulations! We did not find any flagged words in your document!')
            save_name = f"output/{file_name_no_suffix} - Passed.docx"

            
    
     
        
    file_path = f'/tmp/{file_name}'
    doc.save(file_path)
    print('success!')
    
    s3.upload_file(file_path, bucket, save_name)
