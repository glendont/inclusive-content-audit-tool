
import json
import boto3
import os

from docx import Document

def getJobResults(jobId):
    print('getting job results')
    
    pages = []
    textract = boto3.client('textract')
    response = textract.get_document_text_detection(JobId=jobId)
    pages.append(response)

    next_token = None
    if('NextToken' in response):
        next_token = response['NextToken']

    while(next_token):

        response = textract.get_document_text_detection(JobId=jobId, NextToken=next_token)

        pages.append(response)
        next_token = None
        if('NextToken' in response):
            next_token = response['NextToken']
    print('processing done')
    return pages

def lambda_handler(event, context):
    print(event)
    notification_message = json.loads(json.dumps(event))['Records'][0]['Sns']['Message']
    notification_message = json.loads(notification_message)
    
    job_status = notification_message['Status']
    job_tag = notification_message['JobTag']
    job_id = notification_message['JobId']
    document_location = json.loads(json.dumps(notification_message['DocumentLocation']))
    
    
    s3_object = document_location['S3ObjectName']
    bucket = document_location['S3Bucket']
    
    print(job_tag + ' : ' + job_status)
    pdf_text = ''
    
    if(job_status == 'SUCCEEDED'):
        response = getJobResults(job_id)
        print("iterating through pages")
        for resultPage in response:
            for item in resultPage["Blocks"]:
                if item["BlockType"] == "LINE":
                    pdf_text += item["Text"] + '\n'
    
    
        file_name_no_suffix = s3_object.split('/')[-1].split('.')[-2]
        doc = Document()
        doc_para = doc.add_paragraph()
        doc_para.add_run(pdf_text)
        
        file_name = file_name_no_suffix + ".docx"
        save_name = f"input/{file_name}"
        
        file_path = f'/tmp/{file_name}'
        doc.save(file_path)
        print("results successfully saved")
        
        s3 = boto3.client('s3')
        s3.upload_file(file_path, bucket, save_name)
